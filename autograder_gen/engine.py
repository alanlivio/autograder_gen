"""
Autograder generator for creating Gradescope autograder.zip files.
Uses Jinja2 templates and gradescope-utils for proper test generation.
"""

import os
import shutil
import zipfile
import tempfile
from typing import Optional, List, Dict, Any
import yaml
import json
import re
from types import SimpleNamespace
from pathlib import Path
from io import BytesIO, StringIO
import csv
from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape
from autograder_gen.config import Config
from autograder_gen.utils import print_error, print_success, print_warning


class Engine:
    """Generates Gradescope autograder packages from configuration using Jinja templates."""

    def __init__(
        self, config: Config, original_config_dict: Optional[dict] = None
    ):
        """
        Initialize generator with configuration.
        Args:
            config: Config object
            original_config_dict: Optional dictionary of original unparsed YAML configuration
        """
        self.config = config
        self.original_config_dict = (
            original_config_dict  # Store the original JSON config
        )
        self.temp_dir: Optional[Path] = None
        self.templates_dir = Path(__file__).parent / "templates"
        # Set up Jinja environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=select_autoescape(["html", "xml"]),
            trim_blocks=True,
            lstrip_blocks=True,
        )

    def generate(self, output_dir: str) -> str:
        """Generate the autograder.zip file using Jinja templates."""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        # Create temporary directory for autograder files
        self.temp_dir = output_path / "temp_autograder"
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
        self.temp_dir.mkdir()
        # Create source directory within temp
        tests_dir = self.temp_dir / "tests"
        tests_dir.mkdir()
        try:
            self._generate_setup_sh()
            self._generate_run_autograder()
            self._generate_run_tests(tests_dir)
            self._generate_requirements_txt()
            self._generate_metadata_files()
            # Create the zip file
            zip_path = output_path / "autograder.zip"
            self._create_zip(zip_path)
            # Automatically verify generated autograder.zip archive
            verification = self.verify_autograder_zip(zip_path)
            print_success(
                f"Verification for {zip_path.name}: valid={verification['valid']}, entries={verification['total_files']}"
            )
            for err in verification.get("errors", []):
                print_error(f"  [ERROR] {err}")
            for warn in verification.get("warnings", []):
                print_warning(f"  [WARNING] {warn}")
            docx_buffer = self.generate_description_docx()
            with open(output_path / "description.docx", "wb") as f:
                f.write(docx_buffer.getbuffer())
            md_buffer = self.generate_description_md()
            with open(output_path / "description.md", "wb") as f:
                f.write(md_buffer.getbuffer())
            rubric_buffer = self.generate_rubric_csv()
            with open(output_path / "rubric.csv", "wb") as f:
                f.write(rubric_buffer.getbuffer())
            correct_buffer = self.generate_correct_answer_zip()
            with open(output_path / "correct_answer.zip", "wb") as f:
                f.write(correct_buffer.getbuffer())
            wrong_buffer = self.generate_wrong_answer_zip()
            with open(output_path / "wrong_answer.zip", "wb") as f:
                f.write(wrong_buffer.getbuffer())
            return str(zip_path)
        finally:
            # Clean up temporary directory
            if self.temp_dir and self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)

    def verify_autograder_zip(self, zip_source: Any = None) -> Dict[str, Any]:
        """Verify structure and validity of an autograder.zip archive."""
        errors = []
        warnings = []
        checks = []
        try:
            if zip_source is None:
                if self.temp_dir and self.temp_dir.exists():
                    names = [
                        str(f.relative_to(self.temp_dir))
                        for f in self.temp_dir.rglob("*")
                        if f.is_file()
                    ]
                else:
                    return {
                        "valid": False,
                        "checks": [],
                        "errors": ["No zip source or temp directory found to verify"],
                        "warnings": [],
                        "total_files": 0,
                    }
            elif isinstance(zip_source, (str, Path)):
                with zipfile.ZipFile(zip_source, "r") as zf:
                    names = zf.namelist()
            elif isinstance(zip_source, BytesIO):
                with zipfile.ZipFile(zip_source, "r") as zf:
                    names = zf.namelist()
            elif isinstance(zip_source, bytes):
                with zipfile.ZipFile(BytesIO(zip_source), "r") as zf:
                    names = zf.namelist()
            else:
                names = zip_source.namelist() if hasattr(zip_source, "namelist") else []
            if "run_autograder" in names:
                checks.append("Found required executable script: run_autograder")
            else:
                errors.append("Missing required root script: run_autograder")
            if "setup.sh" in names:
                checks.append("Found environment setup script: setup.sh")
            if any(
                n.startswith("source/") or n.startswith("autograder/") or "/" in n
                for n in names
            ):
                checks.append("Archive contains valid folder structures")
            valid = len(errors) == 0
        except Exception as e:
            valid = False
            errors.append(f"Invalid ZIP archive format: {str(e)}")
            names = []
        return {
            "valid": valid,
            "checks": checks,
            "errors": errors,
            "warnings": warnings,
            "total_files": len(names),
        }

    def generate_description_docx(self) -> BytesIO:
        """Generate a Word document containing the assessment description."""
        doc = Document()
        doc.add_heading("Assessment Description", 0)
        for i, question in enumerate(self.config.questions, 1):
            doc.add_heading(f"Question {i}: {question.name}", level=1)
            if hasattr(question, "description") and question.description:
                doc.add_paragraph(question.description)
            # Calculate question total points (excluding hidden file checks)
            question_points = sum(
                item.total_mark
                for item in question.marking_items
                if item.type != "file_exists"
            )
            p = doc.add_paragraph()
            run = p.add_run(f"Total Points: {question_points}")
            run.bold = True
            doc.add_heading("Marking Items", level=2)
            visible_item_idx = 1
            for item in question.marking_items:
                if item.type == "file_exists":
                    continue
                item_name = (
                    getattr(item, "name", "") or f"Marking Item {visible_item_idx}"
                )
                doc.add_heading(f"{visible_item_idx}. {item_name}", level=3)
                doc.add_paragraph(f"Points: {item.total_mark}")
                if item.type == "output_comparison":
                    doc.add_paragraph(
                        f"Requirement: Program must produce specific output for given input in '{item.target_file}'."
                    )
                    if item.expected_input:
                        doc.add_heading("Example Input:", level=4)
                        p = doc.add_paragraph()
                        run = p.add_run(item.expected_input)
                        run.font.name = "Courier New"
                    if item.expected_output:
                        doc.add_heading("Expected Output:", level=4)
                        p = doc.add_paragraph()
                        run = p.add_run(item.expected_output)
                        run.font.name = "Courier New"
                elif item.type == "signature_check":
                    doc.add_paragraph(
                        f"Requirement: Function '{item.function_name}' in '{item.target_file}' must have correct signature."
                    )
                elif item.type == "function_test":
                    doc.add_paragraph(
                        f"Requirement: Function '{item.function_name}' in '{item.target_file}' must pass unit tests."
                    )
                visible_item_idx += 1
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_description_md(self) -> BytesIO:
        lines = ["# Assessment Description", ""]
        for i, question in enumerate(self.config.questions, 1):
            lines.append(f"## Question {i}: {question.name}")
            lines.append("")
            if hasattr(question, "description") and question.description:
                lines.append(question.description)
                lines.append("")
            question_points = sum(
                item.total_mark
                for item in question.marking_items
                if item.type != "file_exists"
            )
            lines.append(f"**Total Points:** {question_points}")
            lines.append("")
            lines.append("### Marking Items")
            lines.append("")
            visible_item_idx = 1
            for item in question.marking_items:
                if item.type == "file_exists":
                    continue
                item_name = (
                    getattr(item, "name", "") or f"Marking Item {visible_item_idx}"
                )
                lines.append(f"#### {visible_item_idx}. {item_name}")
                lines.append("")
                lines.append(f"- **Points:** {item.total_mark}")
                if item.type == "output_comparison":
                    lines.append(
                        f"- **Requirement:** Program must produce specific output for given input in `{item.target_file}`."
                    )
                    if item.expected_input:
                        lines.append("")
                        lines.append("##### Example Input:")
                        lines.append("```")
                        lines.append(item.expected_input)
                        lines.append("```")
                    if item.expected_output:
                        lines.append("")
                        lines.append("##### Expected Output:")
                        lines.append("```")
                        lines.append(item.expected_output)
                        lines.append("```")
                elif item.type == "signature_check":
                    lines.append(
                        f"- **Requirement:** Function `{item.function_name}` in `{item.target_file}` must have correct signature."
                    )
                elif item.type == "function_test":
                    lines.append(
                        f"- **Requirement:** Function `{item.function_name}` in `{item.target_file}` must pass unit tests."
                    )
                lines.append("")
                visible_item_idx += 1
        md_content = "\n".join(lines).strip() + "\n"
        buffer = BytesIO(md_content.encode("utf-8"))
        buffer.seek(0)
        return buffer

    def generate_description_html(self) -> BytesIO:
        summary = self.config.get_config_summary()
        html_lines = [
            "<!DOCTYPE html>",
            '<html lang="en">',
            "<head>",
            '  <meta charset="UTF-8">',
            '  <meta name="viewport" content="width=device-width, initial-scale=1.0">',
            f"  <title>Assessment Description - {self.config.language.capitalize()}</title>",
            "  <style>",
            "    body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; color: #1e293b; max-width: 900px; margin: 0 auto; padding: 2rem; background: #f8fafc; }",
            "    .header { background: #ffffff; padding: 2rem; border-radius: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); margin-bottom: 2rem; border-top: 4px solid #475569; }",
            "    h1 { margin-top: 0; color: #0f172a; }",
            "    .meta-badge { display: inline-block; background: #e2e8f0; color: #334155; padding: 0.25rem 0.75rem; border-radius: 9999px; font-weight: 600; font-size: 0.875rem; margin-right: 0.5rem; }",
            "    .question-card { background: #ffffff; padding: 1.5rem; border-radius: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); margin-bottom: 1.5rem; }",
            "    .question-title { color: #1e293b; margin-top: 0; display: flex; justify-content: space-between; align-items: center; }",
            "    .points-tag { background: #f1f5f9; border: 1px solid #cbd5e1; color: #475569; padding: 0.2rem 0.6rem; border-radius: 4px; font-weight: 600; font-size: 0.875rem; }",
            "    .item-block { border-left: 3px solid #94a3b8; padding-left: 1rem; margin: 1rem 0; }",
            "    pre { background: #0f172a; color: #f8fafc; padding: 1rem; border-radius: 6px; overflow-x: auto; font-family: 'Fira Code', 'Courier New', monospace; font-size: 0.9rem; }",
            "    code { font-family: 'Fira Code', 'Courier New', monospace; background: #e2e8f0; padding: 0.2rem 0.4rem; border-radius: 4px; font-size: 0.9em; }",
            "  </style>",
            "</head>",
            "<body>",
            '  <div class="header">',
            "    <h1>Assessment Description</h1>",
            "    <div>",
            f'      <span class="meta-badge">Language: {summary["language"].capitalize()}</span>',
            f'      <span class="meta-badge">Total Marks: {summary["total_marks"]}</span>',
            f'      <span class="meta-badge">Questions: {summary["total_questions"]}</span>',
            f'      <span class="meta-badge">Time Limit: {summary["global_time_limit"]}s</span>',
            "    </div>",
            "  </div>",
        ]
        for i, question in enumerate(self.config.questions, 1):
            question_points = sum(
                item.total_mark
                for item in question.marking_items
                if item.type != "file_exists"
            )
            html_lines.append('  <div class="question-card">')
            html_lines.append('    <div class="question-title">')
            html_lines.append(f"      <h2>Question {i}: {question.name}</h2>")
            html_lines.append(
                f'      <span class="points-tag">{question_points} pts</span>'
            )
            html_lines.append("    </div>")
            if hasattr(question, "description") and question.description:
                html_lines.append(f"    <p>{question.description}</p>")
            html_lines.append("    <h3>Marking Items</h3>")
            visible_item_idx = 1
            for item in question.marking_items:
                if item.type == "file_exists":
                    continue
                item_name = (
                    getattr(item, "name", "") or f"Marking Item {visible_item_idx}"
                )
                html_lines.append('    <div class="item-block">')
                html_lines.append(
                    f"      <h4>{visible_item_idx}. {item_name} ({item.total_mark} pts)</h4>"
                )
                if item.type == "output_comparison":
                    html_lines.append(
                        f"      <p><strong>Requirement:</strong> Program must produce specific output for target file <code>{item.target_file}</code>.</p>"
                    )
                    if item.expected_input:
                        html_lines.append(
                            "      <p><strong>Example Input:</strong></p>"
                        )
                        html_lines.append(f"      <pre>{item.expected_input}</pre>")
                    if item.expected_output:
                        html_lines.append(
                            "      <p><strong>Expected Output:</strong></p>"
                        )
                        html_lines.append(f"      <pre>{item.expected_output}</pre>")
                elif item.type == "signature_check":
                    html_lines.append(
                        f"      <p><strong>Requirement:</strong> Function <code>{item.function_name}</code> in <code>{item.target_file}</code> must match signature.</p>"
                    )
                elif item.type == "function_test":
                    html_lines.append(
                        f"      <p><strong>Requirement:</strong> Function <code>{item.function_name}</code> in <code>{item.target_file}</code> must pass unit tests.</p>"
                    )
                html_lines.append("    </div>")
                visible_item_idx += 1
            html_lines.append("  </div>")
        html_lines.extend(["</body>", "</html>"])
        html_content = "\n".join(html_lines) + "\n"
        buffer = BytesIO(html_content.encode("utf-8"))
        buffer.seek(0)
        return buffer

    def generate_correct_answer_zip(self) -> BytesIO:
        """Generate a ZIP file with correct implementation skeletons."""
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for filename in self.config.files_necessary:
                content = self._generate_skeleton_content(filename, correct=True)
                zipf.writestr(filename, content)
        buffer.seek(0)
        return buffer

    def generate_wrong_answer_zip(self) -> BytesIO:
        """Generate a ZIP file with incorrect implementation skeletons."""
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for filename in self.config.files_necessary:
                content = self._generate_skeleton_content(filename, correct=False)
                zipf.writestr(filename, content)
        buffer.seek(0)
        return buffer

    def generate_rubric_md(self) -> BytesIO:
        """Generate a Markdown document containing the assessment grading rubric matrix."""
        summary = self.config.get_config_summary()
        lines = [
            "# Assessment Grading Rubric Matrix",
            "",
            "## Overview",
            f"- **Language:** {summary['language']}",
            f"- **Global Time Limit:** {summary['global_time_limit']}s",
            f"- **Total Questions:** {summary['total_questions']}",
            f"- **Total Marking Items:** {summary['total_marking_items']}",
            f"- **Total Marks:** {summary['total_marks']}",
            "",
            "## Detailed Rubric Matrix",
            "",
            "| Question | Marking Item | Type | Target File | Marks | Time Limit | Visibility |",
            "| --- | --- | --- | --- | --- | --- | --- |",
        ]
        for q in self.config.questions:
            for idx, item in enumerate(q.marking_items, 1):
                item_name = getattr(item, "name", "") or f"Item {idx}"
                lines.append(
                    f"| {q.name} | {item_name} | {item.type} | {item.target_file} | {item.total_mark} | {item.time_limit}s | {item.visibility} |"
                )
        lines.append("")
        content = "\n".join(lines)
        buffer = BytesIO(content.encode("utf-8"))
        buffer.seek(0)
        return buffer

    def generate_rubric_csv(self) -> BytesIO:
        """Generate a CSV document containing the assessment grading rubric matrix."""
        s_io = StringIO()
        writer = csv.writer(s_io)
        writer.writerow(
            [
                "Question",
                "Marking Item",
                "Type",
                "Target File",
                "Marks",
                "Time Limit",
                "Visibility",
            ]
        )
        for q in self.config.questions:
            for idx, item in enumerate(q.marking_items, 1):
                item_name = getattr(item, "name", "") or f"Item {idx}"
                writer.writerow(
                    [
                        q.name,
                        item_name,
                        item.type,
                        item.target_file,
                        item.total_mark,
                        item.time_limit,
                        item.visibility,
                    ]
                )
        content = s_io.getvalue()
        buffer = BytesIO(content.encode("utf-8"))
        buffer.seek(0)
        return buffer

    def _generate_skeleton_content(self, target_file: str, correct: bool = True) -> str:
        import ast

        def get_python_literal_str(expected_str: str) -> str:
            try:
                val = ast.literal_eval(expected_str)
                return repr(val)
            except Exception:
                return repr(expected_str)

        def get_java_literal_str(expected_str: str) -> str:
            val_strip = expected_str.strip()
            if val_strip == "True":
                return "true"
            if val_strip == "False":
                return "false"
            if val_strip == "None":
                return "null"
            try:
                int(val_strip)
                return val_strip
            except ValueError:
                pass
            try:
                float(val_strip)
                return val_strip
            except ValueError:
                pass
            escaped = expected_str.replace("\\", "\\\\").replace('"', '\\"')
            return f'"{escaped}"'

        if self.config.language == "python":
            lines = ["# Skeleton for " + target_file, ""]
            functions = set()
            for q in self.config.questions:
                for item in q.marking_items:
                    if item.target_file == target_file:
                        if hasattr(item, "function_name") and item.function_name:
                            functions.add(item.function_name)
            for func in sorted(functions):
                lines.append(f"def {func}(*args, **kwargs):")
                if correct:
                    cases = []
                    for q in self.config.questions:
                        for item in q.marking_items:
                            if (
                                item.target_file == target_file
                                and getattr(item, "function_name", "") == func
                                and item.type == "function_test"
                            ):
                                if hasattr(item, "test_cases") and item.test_cases:
                                    for tc in item.test_cases:
                                        cases.append(tc)
                    if cases:
                        for tc in cases:
                            conds = []
                            args_list = tc.get("args", []) or []
                            kwargs_dict = tc.get("kwargs", {}) or {}
                            if args_list:
                                conds.append(f"args == {tuple(args_list)}")
                            else:
                                conds.append("not args")
                            if kwargs_dict:
                                conds.append(f"kwargs == {kwargs_dict}")
                            else:
                                conds.append("not kwargs")
                            cond_str = " and ".join(conds)
                            expected_val = tc.get("expected", "")
                            expr = get_python_literal_str(expected_val)
                            lines.append(f"    if {cond_str}:")
                            lines.append(f"        return {expr}")
                        lines.append("    pass")
                    else:
                        lines.append("    pass")
                else:
                    lines.append("    return None")
                lines.append("")
            if not functions:
                lines.append("# No specific functions defined for this file.")
                if not correct:
                    lines.append(
                        "# This file might be intentionally wrong or missing logic."
                    )
            return "\n".join(lines)
        elif self.config.language == "java":
            class_name = target_file.replace(".java", "")
            lines = [f"public class {class_name} {{", ""]
            functions = set()
            for q in self.config.questions:
                for item in q.marking_items:
                    if item.target_file == target_file:
                        if hasattr(item, "function_name") and item.function_name:
                            functions.add(item.function_name)
            for func in sorted(functions):
                lines.append(f"    public static double {func}(double a, double b) {{")
                if correct:
                    lines.append("        return a + b;")
                else:
                    lines.append("        return 0.0;")
                lines.append("    }")
                lines.append("")
            lines.append("}")
            return "\n".join(lines)

        return "# Skeleton for " + target_file

    def _generate_setup_sh(self):
        """Generate setup.sh using Jinja template."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"

        template = self.jinja_env.get_template("setup.sh.j2")
        content = template.render(config=self.config)

        setup_file = self.temp_dir / "setup.sh"
        with open(setup_file, "w", encoding="utf-8") as f:
            f.write(content)

        # Make setup.sh executable
        os.chmod(setup_file, 0o755)

    def _generate_run_autograder(self):
        """Generate run_autograder using Jinja template."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"
        template = self.jinja_env.get_template("run_autograder.j2")
        content = template.render(config=self.config)

        run_autograder_file = self.temp_dir / "run_autograder"
        with open(run_autograder_file, "w", encoding="utf-8") as f:
            f.write(content)

        # Make run_autograder executable
        os.chmod(run_autograder_file, 0o755)

    def _generate_run_tests(self, tests_dir: Path):
        """Generate modular test files: main run_tests.py and individual question test files."""
        # Generate main test runner
        assert self.temp_dir is not None, "temp_dir must be set before generating files"
        template = self.jinja_env.get_template("run_tests.py.j2")
        content = template.render(config=self.config)

        run_tests_file = self.temp_dir / "run_tests.py"
        with open(run_tests_file, "w", encoding="utf-8") as f:
            f.write(content)

        # Generate individual question test files
        self._generate_question_test_files(tests_dir)

    def _generate_question_test_files(self, tests_dir: Path):
        """Generate individual test files for each question."""
        question_template = self.jinja_env.get_template("test_question.py.j2")

        for idx, question in enumerate(self.config.questions, 1):
            # Use question number for filename
            question_filename = f"question_{idx}"

            # Preprocess marking items to ensure output comparison tests have proper newlines
            processed_question = self._preprocess_question_for_output_comparison(
                question
            )

            # Generate content for this question
            content = question_template.render(
                config=self.config, question=processed_question, question_number=idx
            )

            # Write the question test file
            test_file = tests_dir / f"{question_filename}_test.py"
            with open(test_file, "w", encoding="utf-8") as f:
                f.write(content)

    def _preprocess_question_for_output_comparison(self, question):
        """Preprocess question to add newlines to expected output for output comparison tests."""
        # Create a copy of the question with processed marking items

        processed_question = SimpleNamespace()
        processed_question.name = question.name
        processed_question.marking_items = []

        for item in question.marking_items:
            processed_item = SimpleNamespace()
            # Copy all attributes from the original item
            for attr in dir(item):
                if not attr.startswith("_") and not attr.startswith("model_"):
                    setattr(processed_item, attr, getattr(item, attr))

            # For Python output comparison tests, ensure expected_output has a newline if it doesn't end with one
            # This matches Python's print() behavior which automatically adds newlines
            if (
                self.config.language == "python"
                and hasattr(processed_item, "type")
                and processed_item.type == "output_comparison"
                and hasattr(processed_item, "expected_output")
                and processed_item.expected_output
                and not processed_item.expected_output.endswith("\n")
            ):
                processed_item.expected_output += "\n"

            processed_question.marking_items.append(processed_item)

        return processed_question

    def _sanitize_filename(self, name: str) -> str:
        """Convert question name to a safe Python module filename."""
        # Convert to lowercase and replace problematic characters
        safe_name = name.lower()
        safe_name = safe_name.replace(" ", "_")
        safe_name = safe_name.replace("-", "_")
        safe_name = safe_name.replace(".", "_")
        safe_name = safe_name.replace("(", "")
        safe_name = safe_name.replace(")", "")
        safe_name = safe_name.replace("[", "")
        safe_name = safe_name.replace("]", "")
        safe_name = safe_name.replace("/", "_")
        safe_name = safe_name.replace("\\", "_")
        safe_name = safe_name.replace(":", "_")
        safe_name = safe_name.replace(";", "_")
        safe_name = safe_name.replace(",", "_")
        safe_name = safe_name.replace("?", "_")
        safe_name = safe_name.replace("!", "_")
        safe_name = safe_name.replace("@", "_")
        safe_name = safe_name.replace("#", "_")
        safe_name = safe_name.replace("$", "_")
        safe_name = safe_name.replace("%", "_")
        safe_name = safe_name.replace("^", "_")
        safe_name = safe_name.replace("&", "_")
        safe_name = safe_name.replace("*", "_")
        safe_name = safe_name.replace("+", "_")
        safe_name = safe_name.replace("=", "_")
        safe_name = safe_name.replace("|", "_")
        safe_name = safe_name.replace("<", "_")
        safe_name = safe_name.replace(">", "_")

        # Remove multiple consecutive underscores
        safe_name = re.sub(r"_+", "_", safe_name)

        # Remove leading/trailing underscores
        safe_name = safe_name.strip("_")

        # Ensure it's a valid Python identifier
        if not safe_name or safe_name[0].isdigit():
            safe_name = "question_" + safe_name

        return safe_name

    def _generate_requirements_txt(self):
        """Generate requirements.txt using Jinja template."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"

        template = self.jinja_env.get_template("requirements.txt.j2")
        content = template.render(config=self.config)

        requirements_file = self.temp_dir / "requirements.txt"
        with open(requirements_file, "w", encoding="utf-8") as f:
            f.write(content)

    def _generate_metadata_files(self):
        """Generate metadata and configuration files."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"

        # Save the original configuration if provided
        if self.original_config_dict:
            original_config_file = self.temp_dir / "autograder_gen.yaml"
            with open(original_config_file, "w", encoding="utf-8") as f:
                yaml.dump(
                    self.original_config_dict,
                    f,
                    default_flow_style=False,
                    sort_keys=False,
                )

        # Create a README for the autograder
        readme_content = f"""# Autograder Package

Generated by TIF Autograder Tool

## Configuration Summary
- **Language**: {self.config.language}
- **Questions**: {len(self.config.questions)}
- **Total Marking Items**: {sum(len(q.marking_items) for q in self.config.questions)}
- **Total Points**: {sum(sum(item.total_mark for item in q.marking_items) for q in self.config.questions)}
- **Required Files**: {', '.join(self.config.files_necessary) if self.config.files_necessary else 'None specified'}

## Package Structure
```
autograder.zip
├── setup.sh                 # Environment setup script
├── run_autograder          # Main autograder execution script
├── run_tests.py            # Primary test runner using gradescope-utils
├── requirements.txt        # Python dependencies
├── tests/                  # Individual test files for each question
│   ├── question_1_test.py
│   ├── question_2_test.py
│   └── ...
├── autograder_gen.yaml     # Original configuration file
└── README.md              # This file
```

## Test Types Supported
- **file_exists**: Checks if required files are present in submission
- **output_comparison**: Compares program output with expected results
- **signature_check**: Validates function signatures and parameters
- **function_test**: Tests function behavior with specific inputs and expected outputs

## Global Settings
- **Global Time Limit**: {getattr(self.config, 'global_time_limit', 'Not set')} seconds
- **Points Precision**: {getattr(self.config, 'points_precision', 1)} decimal place(s)

## Questions and Marking Items"""

        for i, question in enumerate(self.config.questions, 1):
            readme_content += f"\n\n### Question {i}: {question.name}\n"

            # Calculate question total points
            question_points = sum(item.total_mark for item in question.marking_items)
            readme_content += f"**Total Points**: {question_points}\n\n"

            for j, item in enumerate(question.marking_items, 1):
                item_name = getattr(item, "name", "") or f"Marking Item {j}"
                readme_content += f"#### {j}. {item_name}\n"
                readme_content += f"- **Type**: {item.type.replace('_', ' ').title()}\n"
                readme_content += f"- **Target File**: {item.target_file}\n"
                readme_content += f"- **Points**: {item.total_mark}\n"

                # Add time limit if specified
                if hasattr(item, "time_limit") and item.time_limit:
                    readme_content += f"- **Time Limit**: {item.time_limit} seconds\n"

                # Add visibility setting
                if hasattr(item, "visibility") and item.visibility:
                    visibility_map = {
                        "hidden": "Hidden from students",
                        "visible": "Visible to students immediately",
                        "after_due_date": "Visible after due date",
                        "after_published": "Visible after grades published",
                    }
                    readme_content += f"- **Visibility**: {visibility_map.get(item.visibility, item.visibility)}\n"

                # Add type-specific details
                if item.type == "function_test":
                    if hasattr(item, "function_name") and item.function_name:
                        readme_content += f"- **Function**: `{item.function_name}()`\n"
                    if hasattr(item, "test_cases") and item.test_cases:
                        readme_content += (
                            f"- **Test Cases**: {len(item.test_cases)} case(s)\n"
                        )

                elif item.type == "signature_check":
                    if hasattr(item, "function_name") and item.function_name:
                        readme_content += f"- **Function**: `{item.function_name}()`\n"
                    if (
                        hasattr(item, "expected_parameters")
                        and item.expected_parameters
                    ):
                        readme_content += (
                            f"- **Expected Parameters**: `{item.expected_parameters}`\n"
                        )

                elif item.type == "output_comparison":
                    if hasattr(item, "expected_input") and item.expected_input:
                        input_lines = item.expected_input.count("\n") + 1
                        readme_content += f"- **Input Lines**: {input_lines}\n"
                    if hasattr(item, "expected_output") and item.expected_output:
                        output_lines = item.expected_output.count("\n") + 1
                        readme_content += (
                            f"- **Expected Output Lines**: {output_lines}\n"
                        )

                readme_content += "\n"

        # Add execution information
        readme_content += f"""
## Execution Details

### Setup Process
1. **Environment Setup**: `setup.sh` installs required packages and prepares the testing environment
2. **Test Execution**: `run_autograder` executes `run_tests.py` which runs all question test files
3. **Results Collection**: Results are formatted using gradescope-utils and written to `/autograder/results/results.json`

### File Requirements
Students must submit the following files:
"""

        if self.config.files_necessary:
            for file in self.config.files_necessary:
                readme_content += f"- `{file}`\n"
        else:
            readme_content += (
                "- No specific files required (will be determined by marking items)\n"
            )

        readme_content += f"""
### Points Distribution
"""

        # Create points breakdown by question
        for i, question in enumerate(self.config.questions, 1):
            question_points = sum(item.total_mark for item in question.marking_items)
            readme_content += f"- **Question {i}**: {question_points} points\n"

        total_points = sum(
            sum(item.total_mark for item in q.marking_items)
            for q in self.config.questions
        )
        readme_content += f"- **Total Possible**: {total_points} points\n"

        readme_content += f"""
## Technical Notes

- Generated using TIF Autograder Tool
- Uses gradescope-utils for test framework compatibility
- Supports Python {self.config.language} submissions
- All tests run in isolated environments with proper timeout handling
- Results are automatically formatted for Gradescope integration

For questions about this autograder configuration, refer to the original `autograder_gen.yaml` file included in this package.
"""

        readme_file = self.temp_dir / "README.md"
        with open(readme_file, "w", encoding="utf-8") as f:
            f.write(readme_content)

    def _create_zip(self, zip_path: Path):
        """Create the autograder.zip file with proper structure."""
        assert self.temp_dir is not None, "temp_dir must be set before creating zip"

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file_path in self.temp_dir.rglob("*"):
                if file_path.is_file():
                    arcname = file_path.relative_to(self.temp_dir)
                    zipf.write(file_path, arcname)

