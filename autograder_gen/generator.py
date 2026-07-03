"""
Autograder generator for creating Gradescope autograder.zip files.
Uses Jinja2 templates and gradescope-utils for proper test generation.
"""

import os
import shutil
import zipfile
from typing import Optional, List, Dict, Any
import yaml
import re
from types import SimpleNamespace
from pathlib import Path

from io import BytesIO
from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape
from autograder_gen.config import AutograderConfig


class AutograderGenerator:
    """Generates Gradescope autograder packages from configuration using Jinja templates."""

    def __init__(
        self, config: AutograderConfig, original_config_dict: Optional[dict] = None
    ):
        self.config = config
        self.original_config_dict = (
            original_config_dict  # Store the original JSON config
        )
        self.temp_dir: Optional[Path] = None
        self.templates_dir = Path(__file__).parent / "templates"

        # Set up Jinja environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=select_autoescape(["html", "xml"]),
            trim_blocks=True,
            lstrip_blocks=True,
        )

    def generate(self, output_dir: str) -> str:
        """Generate the autograder.zip file using Jinja templates."""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        # Create temporary directory for autograder files
        self.temp_dir = output_path / "temp_autograder"
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
        self.temp_dir.mkdir()

        # Create source directory within temp
        tests_dir = self.temp_dir / "tests"
        tests_dir.mkdir()

        try:
            # Generate all autograder files using templates
            self._generate_setup_sh()
            self._generate_run_autograder()
            self._generate_run_tests(tests_dir)
            self._generate_requirements_txt()
            self._generate_metadata_files()

            # Create the zip file
            zip_path = output_path / "autograder.zip"
            self._create_zip(zip_path)

            return str(zip_path)

        finally:
            # Clean up temporary directory
            if self.temp_dir and self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)

    def generate_description_docx(self) -> BytesIO:
        """Generate a Word document containing the assessment description."""
        doc = Document()
        doc.add_heading("Assessment Description", 0)

        for i, question in enumerate(self.config.questions, 1):
            doc.add_heading(f"Question {i}: {question.name}", level=1)
            if hasattr(question, "description") and question.description:
                doc.add_paragraph(question.description)

            # Calculate question total points (excluding hidden file checks)
            question_points = sum(
                item.total_mark
                for item in question.marking_items
                if item.type != "file_exists"
            )
            p = doc.add_paragraph()
            run = p.add_run(f"Total Points: {question_points}")
            run.bold = True

            doc.add_heading("Marking Items", level=2)
            visible_item_idx = 1
            for item in question.marking_items:
                if item.type == "file_exists":
                    continue

                item_name = getattr(item, "name", "") or f"Marking Item {visible_item_idx}"
                doc.add_heading(f"{visible_item_idx}. {item_name}", level=3)
                doc.add_paragraph(f"Points: {item.total_mark}")
                
                if item.type == "output_comparison":
                    doc.add_paragraph(f"Requirement: Program must produce specific output for given input in '{item.target_file}'.")
                    if item.expected_input:
                        doc.add_heading("Example Input:", level=4)
                        p = doc.add_paragraph()
                        run = p.add_run(item.expected_input)
                        run.font.name = 'Courier New'
                    
                    if item.expected_output:
                        doc.add_heading("Expected Output:", level=4)
                        p = doc.add_paragraph()
                        run = p.add_run(item.expected_output)
                        run.font.name = 'Courier New'
                elif item.type == "signature_check":
                    doc.add_paragraph(f"Requirement: Function '{item.function_name}' in '{item.target_file}' must have correct signature.")
                elif item.type == "function_test":
                    doc.add_paragraph(f"Requirement: Function '{item.function_name}' in '{item.target_file}' must pass unit tests.")
                
                visible_item_idx += 1

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_correct_answer_zip(self) -> BytesIO:
        """Generate a ZIP file with correct implementation skeletons."""
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for filename in self.config.files_necessary:
                content = self._generate_skeleton_content(filename, correct=True)
                zipf.writestr(filename, content)
        buffer.seek(0)
        return buffer

    def generate_wrong_answer_zip(self) -> BytesIO:
        """Generate a ZIP file with incorrect implementation skeletons."""
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for filename in self.config.files_necessary:
                content = self._generate_skeleton_content(filename, correct=False)
                zipf.writestr(filename, content)
        buffer.seek(0)
        return buffer

    def _generate_skeleton_content(self, target_file: str, correct: bool = True) -> str:
        """Generate skeleton code for a given file."""
        if self.config.language == "python":
            lines = ["# Skeleton for " + target_file, ""]
            
            # Find all functions related to this file
            functions = set()
            for q in self.config.questions:
                for item in q.marking_items:
                    if item.target_file == target_file:
                        if hasattr(item, "function_name") and item.function_name:
                            functions.add(item.function_name)
            
            for func in sorted(functions):
                lines.append(f"def {func}(*args, **kwargs):")
                if correct:
                    lines.append("    # TODO: Implement correct logic")
                    lines.append("    pass")
                else:
                    lines.append("    # TODO: Implement incorrect logic for testing")
                    lines.append("    return None")
                lines.append("")
            
            if not functions:
                lines.append("# No specific functions defined for this file.")
                if not correct:
                    lines.append("# This file might be intentionally wrong or missing logic.")

            return "\n".join(lines)
        
        elif self.config.language == "java":
            # Very basic Java skeleton
            class_name = target_file.replace(".java", "")
            lines = [f"public class {class_name} {{", ""]
            
            functions = set()
            for q in self.config.questions:
                for item in q.marking_items:
                    if item.target_file == target_file:
                        if hasattr(item, "function_name") and item.function_name:
                            functions.add(item.function_name)
            
            for func in sorted(functions):
                lines.append(f"    public static Object {func}(Object... args) {{")
                if correct:
                    lines.append("        // TODO: Implement correct logic")
                    lines.append("        return null;")
                else:
                    lines.append("        // TODO: Implement incorrect logic")
                    lines.append("        throw new RuntimeException(\"Not implemented\");")
                lines.append("    }")
                lines.append("")
                
            lines.append("}")
            return "\n".join(lines)

        return "# Skeleton for " + target_file

    def _generate_setup_sh(self):
        """Generate setup.sh using Jinja template."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"

        template = self.jinja_env.get_template("setup.sh.j2")
        content = template.render(config=self.config)

        setup_file = self.temp_dir / "setup.sh"
        with open(setup_file, "w", encoding="utf-8") as f:
            f.write(content)

        # Make setup.sh executable
        os.chmod(setup_file, 0o755)

    def _generate_run_autograder(self):
        """Generate run_autograder using Jinja template."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"
        template = self.jinja_env.get_template("run_autograder.j2")
        content = template.render(config=self.config)

        run_autograder_file = self.temp_dir / "run_autograder"
        with open(run_autograder_file, "w", encoding="utf-8") as f:
            f.write(content)

        # Make run_autograder executable
        os.chmod(run_autograder_file, 0o755)

    def _generate_run_tests(self, tests_dir: Path):
        """Generate modular test files: main run_tests.py and individual question test files."""
        # Generate main test runner
        assert self.temp_dir is not None, "temp_dir must be set before generating files"
        template = self.jinja_env.get_template("run_tests.py.j2")
        content = template.render(config=self.config)

        run_tests_file = self.temp_dir / "run_tests.py"
        with open(run_tests_file, "w", encoding="utf-8") as f:
            f.write(content)

        # Generate individual question test files
        self._generate_question_test_files(tests_dir)

    def _generate_question_test_files(self, tests_dir: Path):
        """Generate individual test files for each question."""
        question_template = self.jinja_env.get_template("test_question.py.j2")

        for idx, question in enumerate(self.config.questions, 1):
            # Use question number for filename
            question_filename = f"question_{idx}"

            # Preprocess marking items to ensure output comparison tests have proper newlines
            processed_question = self._preprocess_question_for_output_comparison(
                question
            )

            # Generate content for this question
            content = question_template.render(
                config=self.config, question=processed_question, question_number=idx
            )

            # Write the question test file
            test_file = tests_dir / f"{question_filename}_test.py"
            with open(test_file, "w", encoding="utf-8") as f:
                f.write(content)

    def _preprocess_question_for_output_comparison(self, question):
        """Preprocess question to add newlines to expected output for output comparison tests."""
        # Create a copy of the question with processed marking items

        processed_question = SimpleNamespace()
        processed_question.name = question.name
        processed_question.marking_items = []

        for item in question.marking_items:
            processed_item = SimpleNamespace()
            # Copy all attributes from the original item
            for attr in dir(item):
                if not attr.startswith("_") and not attr.startswith("model_"):
                    setattr(processed_item, attr, getattr(item, attr))

            # For Python output comparison tests, ensure expected_output has a newline if it doesn't end with one
            # This matches Python's print() behavior which automatically adds newlines
            if (
                self.config.language == "python"
                and hasattr(processed_item, "type")
                and processed_item.type == "output_comparison"
                and hasattr(processed_item, "expected_output")
                and processed_item.expected_output
                and not processed_item.expected_output.endswith("\n")
            ):
                processed_item.expected_output += "\n"

            processed_question.marking_items.append(processed_item)

        return processed_question

    def _sanitize_filename(self, name: str) -> str:
        """Convert question name to a safe Python module filename."""
        # Convert to lowercase and replace problematic characters
        safe_name = name.lower()
        safe_name = safe_name.replace(" ", "_")
        safe_name = safe_name.replace("-", "_")
        safe_name = safe_name.replace(".", "_")
        safe_name = safe_name.replace("(", "")
        safe_name = safe_name.replace(")", "")
        safe_name = safe_name.replace("[", "")
        safe_name = safe_name.replace("]", "")
        safe_name = safe_name.replace("/", "_")
        safe_name = safe_name.replace("\\", "_")
        safe_name = safe_name.replace(":", "_")
        safe_name = safe_name.replace(";", "_")
        safe_name = safe_name.replace(",", "_")
        safe_name = safe_name.replace("?", "_")
        safe_name = safe_name.replace("!", "_")
        safe_name = safe_name.replace("@", "_")
        safe_name = safe_name.replace("#", "_")
        safe_name = safe_name.replace("$", "_")
        safe_name = safe_name.replace("%", "_")
        safe_name = safe_name.replace("^", "_")
        safe_name = safe_name.replace("&", "_")
        safe_name = safe_name.replace("*", "_")
        safe_name = safe_name.replace("+", "_")
        safe_name = safe_name.replace("=", "_")
        safe_name = safe_name.replace("|", "_")
        safe_name = safe_name.replace("<", "_")
        safe_name = safe_name.replace(">", "_")

        # Remove multiple consecutive underscores
        safe_name = re.sub(r"_+", "_", safe_name)

        # Remove leading/trailing underscores
        safe_name = safe_name.strip("_")

        # Ensure it's a valid Python identifier
        if not safe_name or safe_name[0].isdigit():
            safe_name = "question_" + safe_name

        return safe_name

    def _generate_requirements_txt(self):
        """Generate requirements.txt using Jinja template."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"

        template = self.jinja_env.get_template("requirements.txt.j2")
        content = template.render(config=self.config)

        requirements_file = self.temp_dir / "requirements.txt"
        with open(requirements_file, "w", encoding="utf-8") as f:
            f.write(content)

    def _generate_metadata_files(self):
        """Generate metadata and configuration files."""
        assert self.temp_dir is not None, "temp_dir must be set before generating files"

        # Save the original configuration if provided
        if self.original_config_dict:
            original_config_file = self.temp_dir / "autograder_gen.yaml"
            with open(original_config_file, "w", encoding="utf-8") as f:
                yaml.dump(
                    self.original_config_dict,
                    f,
                    default_flow_style=False,
                    sort_keys=False,
                )

        # Create a README for the autograder
        readme_content = f"""# Autograder Package

Generated by TIF Autograder Tool

## Configuration Summary
- **Language**: {self.config.language}
- **Questions**: {len(self.config.questions)}
- **Total Marking Items**: {sum(len(q.marking_items) for q in self.config.questions)}
- **Total Points**: {sum(sum(item.total_mark for item in q.marking_items) for q in self.config.questions)}
- **Required Files**: {', '.join(self.config.files_necessary) if self.config.files_necessary else 'None specified'}

## Package Structure
```
autograder.zip
├── setup.sh                 # Environment setup script
├── run_autograder          # Main autograder execution script
├── run_tests.py            # Primary test runner using gradescope-utils
├── requirements.txt        # Python dependencies
├── tests/                  # Individual test files for each question
│   ├── question_1_test.py
│   ├── question_2_test.py
│   └── ...
├── autograder_gen.yaml     # Original configuration file
└── README.md              # This file
```

## Test Types Supported
- **file_exists**: Checks if required files are present in submission
- **output_comparison**: Compares program output with expected results
- **signature_check**: Validates function signatures and parameters
- **function_test**: Tests function behavior with specific inputs and expected outputs

## Global Settings
- **Global Time Limit**: {getattr(self.config, 'global_time_limit', 'Not set')} seconds
- **Points Precision**: {getattr(self.config, 'points_precision', 1)} decimal place(s)

## Questions and Marking Items"""

        for i, question in enumerate(self.config.questions, 1):
            readme_content += f"\n\n### Question {i}: {question.name}\n"

            # Calculate question total points
            question_points = sum(item.total_mark for item in question.marking_items)
            readme_content += f"**Total Points**: {question_points}\n\n"

            for j, item in enumerate(question.marking_items, 1):
                item_name = getattr(item, "name", "") or f"Marking Item {j}"
                readme_content += f"#### {j}. {item_name}\n"
                readme_content += f"- **Type**: {item.type.replace('_', ' ').title()}\n"
                readme_content += f"- **Target File**: {item.target_file}\n"
                readme_content += f"- **Points**: {item.total_mark}\n"

                # Add time limit if specified
                if hasattr(item, "time_limit") and item.time_limit:
                    readme_content += f"- **Time Limit**: {item.time_limit} seconds\n"

                # Add visibility setting
                if hasattr(item, "visibility") and item.visibility:
                    visibility_map = {
                        "hidden": "Hidden from students",
                        "visible": "Visible to students immediately",
                        "after_due_date": "Visible after due date",
                        "after_published": "Visible after grades published",
                    }
                    readme_content += f"- **Visibility**: {visibility_map.get(item.visibility, item.visibility)}\n"

                # Add type-specific details
                if item.type == "function_test":
                    if hasattr(item, "function_name") and item.function_name:
                        readme_content += f"- **Function**: `{item.function_name}()`\n"
                    if hasattr(item, "test_cases") and item.test_cases:
                        readme_content += (
                            f"- **Test Cases**: {len(item.test_cases)} case(s)\n"
                        )

                elif item.type == "signature_check":
                    if hasattr(item, "function_name") and item.function_name:
                        readme_content += f"- **Function**: `{item.function_name}()`\n"
                    if (
                        hasattr(item, "expected_parameters")
                        and item.expected_parameters
                    ):
                        readme_content += (
                            f"- **Expected Parameters**: `{item.expected_parameters}`\n"
                        )

                elif item.type == "output_comparison":
                    if hasattr(item, "expected_input") and item.expected_input:
                        input_lines = item.expected_input.count("\n") + 1
                        readme_content += f"- **Input Lines**: {input_lines}\n"
                    if hasattr(item, "expected_output") and item.expected_output:
                        output_lines = item.expected_output.count("\n") + 1
                        readme_content += (
                            f"- **Expected Output Lines**: {output_lines}\n"
                        )

                readme_content += "\n"

        # Add execution information
        readme_content += f"""
## Execution Details

### Setup Process
1. **Environment Setup**: `setup.sh` installs required packages and prepares the testing environment
2. **Test Execution**: `run_autograder` executes `run_tests.py` which runs all question test files
3. **Results Collection**: Results are formatted using gradescope-utils and written to `/autograder/results/results.json`

### File Requirements
Students must submit the following files:
"""

        if self.config.files_necessary:
            for file in self.config.files_necessary:
                readme_content += f"- `{file}`\n"
        else:
            readme_content += (
                "- No specific files required (will be determined by marking items)\n"
            )

        readme_content += f"""
### Points Distribution
"""

        # Create points breakdown by question
        for i, question in enumerate(self.config.questions, 1):
            question_points = sum(item.total_mark for item in question.marking_items)
            readme_content += f"- **Question {i}**: {question_points} points\n"

        total_points = sum(
            sum(item.total_mark for item in q.marking_items)
            for q in self.config.questions
        )
        readme_content += f"- **Total Possible**: {total_points} points\n"

        readme_content += f"""
## Technical Notes

- Generated using TIF Autograder Tool
- Uses gradescope-utils for test framework compatibility
- Supports Python {self.config.language} submissions
- All tests run in isolated environments with proper timeout handling
- Results are automatically formatted for Gradescope integration

For questions about this autograder configuration, refer to the original `autograder_gen.yaml` file included in this package.
"""

        readme_file = self.temp_dir / "README.md"
        with open(readme_file, "w", encoding="utf-8") as f:
            f.write(readme_content)

    def _create_zip(self, zip_path: Path):
        """Create the autograder.zip file with proper structure."""
        assert self.temp_dir is not None, "temp_dir must be set before creating zip"

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file_path in self.temp_dir.rglob("*"):
                if file_path.is_file():
                    arcname = file_path.relative_to(self.temp_dir)
                    zipf.write(file_path, arcname)
